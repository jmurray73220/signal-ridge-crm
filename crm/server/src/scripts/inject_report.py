"""
Inject Claude-generated report content into a .docx template.
Preserves the template's styles and formatting.

Usage: python inject_report.py <template.docx> <content.txt> <output.docx>
"""

import sys
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_markdown_to_sections(content):
    """Parse markdown content into structured sections."""
    lines = content.split('\n')
    sections = []
    current_section = None
    current_body = []

    for line in lines:
        # H1
        if line.startswith('# ') and not line.startswith('## '):
            if current_section:
                current_section['body'] = '\n'.join(current_body).strip()
                sections.append(current_section)
            current_section = {'level': 1, 'title': line[2:].strip(), 'body': ''}
            current_body = []
        # H2
        elif line.startswith('## '):
            if current_section:
                current_section['body'] = '\n'.join(current_body).strip()
                sections.append(current_section)
            current_section = {'level': 2, 'title': line[3:].strip(), 'body': ''}
            current_body = []
        # H3
        elif line.startswith('### '):
            if current_section:
                current_section['body'] = '\n'.join(current_body).strip()
                sections.append(current_section)
            current_section = {'level': 3, 'title': line[4:].strip(), 'body': ''}
            current_body = []
        else:
            current_body.append(line)

    if current_section:
        current_section['body'] = '\n'.join(current_body).strip()
        sections.append(current_section)

    # If no markdown headings found, treat entire content as one section
    if not sections:
        sections = [{'level': 1, 'title': 'Report', 'body': content.strip()}]

    return sections


def get_template_styles(doc):
    """Extract style information from the template."""
    styles = {
        'heading1': None,
        'heading2': None,
        'heading3': None,
        'normal': None,
    }

    for style in doc.styles:
        name_lower = style.name.lower() if style.name else ''
        if name_lower == 'heading 1':
            styles['heading1'] = style
        elif name_lower == 'heading 2':
            styles['heading2'] = style
        elif name_lower == 'heading 3':
            styles['heading3'] = style
        elif name_lower == 'normal':
            styles['normal'] = style

    return styles


def add_formatted_paragraph(doc, text, style_name='Normal', bold=False):
    """Add a paragraph with formatting, handling bold markdown markers."""
    para = doc.add_paragraph()
    try:
        para.style = doc.styles[style_name]
    except KeyError:
        pass

    # Split on **bold** markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            run = para.add_run(part)
            if bold:
                run.bold = True

    return para


def inject_content(template_path, content_path, output_path):
    """Inject report content into the template."""
    doc = Document(template_path)
    styles = get_template_styles(doc)

    with open(content_path, 'r', encoding='utf-8') as f:
        content = f.read()

    sections = parse_markdown_to_sections(content)

    # Clear template content after the first page/section
    # Keep the first paragraph if it looks like a cover page element
    paragraphs_to_keep = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()
        if not text or 'template' in text or 'report' in text:
            paragraphs_to_keep = i + 1
        else:
            break

    # Remove all paragraphs after the ones we want to keep
    # We do this by clearing and rebuilding
    while len(doc.paragraphs) > max(paragraphs_to_keep, 1):
        p = doc.paragraphs[-1]._element
        p.getparent().remove(p)

    # Clear the last remaining paragraph if it's a placeholder
    if doc.paragraphs:
        last = doc.paragraphs[-1]
        if 'template' in last.text.lower() or not last.text.strip():
            last.clear()

    # Inject the report content
    for section in sections:
        level = section['level']
        title = section['title']
        body = section['body']

        # Add heading
        if level == 1:
            style_name = 'Heading 1'
        elif level == 2:
            style_name = 'Heading 2'
        else:
            style_name = 'Heading 3'

        try:
            heading = doc.add_heading(title, level=level)
        except Exception:
            heading = doc.add_paragraph(title)
            heading.runs[0].bold = True

        # Add body paragraphs
        if body:
            paragraphs = body.split('\n\n')
            for para_text in paragraphs:
                para_text = para_text.strip()
                if not para_text:
                    continue

                # Handle bullet points
                if para_text.startswith('- ') or para_text.startswith('* '):
                    for bullet_line in para_text.split('\n'):
                        bullet_line = bullet_line.strip()
                        if bullet_line.startswith('- ') or bullet_line.startswith('* '):
                            bullet_line = bullet_line[2:]
                        if bullet_line:
                            try:
                                add_formatted_paragraph(doc, bullet_line, 'List Bullet')
                            except Exception:
                                add_formatted_paragraph(doc, f"• {bullet_line}")
                # Handle numbered lists
                elif re.match(r'^\d+\.', para_text):
                    for num_line in para_text.split('\n'):
                        num_line = num_line.strip()
                        if num_line:
                            add_formatted_paragraph(doc, num_line)
                else:
                    # Regular paragraph - handle inline bold
                    add_formatted_paragraph(doc, para_text)

    doc.save(output_path)


if __name__ == '__main__':
    if len(sys.argv) != 4:
        print(f"Usage: {sys.argv[0]} <template.docx> <content.txt> <output.docx>")
        sys.exit(1)

    inject_content(sys.argv[1], sys.argv[2], sys.argv[3])
